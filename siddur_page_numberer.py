#!/usr/bin/env python3
"""
Siddur Page Numberer v14

Adds accurate page numbers to Siddur Processor Word documents.
Uses LibreOffice to determine actual page breaks.

v14 changes:
  - New Step 7: Cleanup pass with three tidy-up rules applied after numbering
  - Rule 1: Collapse empty suffix pairs (e.g., 1197a+1197b with no content → 1197)
  - Rule 2: Label unlabeled overflow pages (e.g., content after 1195b → 1195c)
  - Rule 3: Consolidate blank page separators (move page break to PN, remove
    empty Normal paragraph that could render as a blank page)

v13 changes:
  - split_paragraph_at_last_period tries earlier periods when last gives empty 2nd half
  - New split_paragraph_at_blank_line for transliteration with \n\n paragraph structure
  - Transliteration split strategy: blank-line first, then 40% midpoint, then fallback
  - Short transliterations (<150 chars) are NOT split — page number inserted before
  - Lowered midpoint threshold from 1/2 to 1/3 (scans bottom two-thirds of page)
  - New Pass 4: finds natural paragraph breaks (period-ending paragraphs) in lower 2/3
  - is_attribution checks paragraph style first; text heuristics only for unstyled paras
  - Pass 2 excludes section headers from "new reading" detection
  - adjust_insert_for_grouping prevents stranding section headers and attributions
  - Pass 2 new-reading handler won't return a section header as the break point

v12 changes:
  - Attribution lines now serve as preferred break points (Pass 3)
  - Instruction/navigation notes (is_instruction_note) kept with preceding reading
  - is_keep_with_above() groups attributions + notes as a unit
  - Pass 1 (section headers) scans back past attributions for cleaner breaks
  - _strip_bidi() applied in paragraph-page mapping for better Hebrew matching
  - _strip_marks() fallback strips Hebrew vowels for PDF text matching
  - Empty Word pages filtered from ranges (eliminates phantom labels)
  - split_paragraph_at_midpoint biased to 65% (default) / 75% (solo pages)
  - find_best_insert_point returns structural flag; Step 5 only repositions
    for structural breaks (headers, new readings, attributions)
  - Default fallback respects keep-with-above on last paragraph

Requirements:
    pip3 install python-docx pypdf lxml
    LibreOffice (https://www.libreoffice.org)

Usage:
    python3 siddur_page_numberer.py my_siddur_processed.docx
"""

import argparse, os, re, subprocess, sys, tempfile, unicodedata
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy


def find_libreoffice():
    # Check for soffice.py wrapper first (Claude sandbox)
    if os.path.exists('/mnt/skills/public/docx/scripts/office/soffice.py'):
        return '/mnt/skills/public/docx/scripts/office/soffice.py'
    for p in [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/local/bin/libreoffice", "/usr/local/bin/soffice",
        "/usr/bin/libreoffice", "/usr/bin/soffice",
        "libreoffice", "soffice",
    ]:
        try:
            if subprocess.run([p, "--version"], capture_output=True, timeout=10).returncode == 0:
                return p
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
    return None


def render_to_pdf_pages(docx_path, lo_bin):
    # Try the soffice.py wrapper first (Claude sandbox environment)
    soffice_wrapper = '/mnt/skills/public/docx/scripts/office/soffice.py'
    pdf_path = None
    basename = Path(docx_path).stem

    if os.path.exists(soffice_wrapper):
        subprocess.run(
            ['python3', soffice_wrapper, '--headless', '--convert-to', 'pdf', str(docx_path)],
            capture_output=True, text=True, timeout=300
        )
        for candidate in [
            f'/{basename}.pdf',
            str(Path(docx_path).with_suffix('.pdf')),
            f'{basename}.pdf',
        ]:
            if os.path.exists(candidate):
                pdf_path = candidate
                break

    if pdf_path is None:
        # Fall back to direct LibreOffice
        with tempfile.TemporaryDirectory() as tmp:
            subprocess.run(
                [lo_bin or "soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, str(docx_path)],
                capture_output=True, text=True, timeout=300, check=True
            )
            pdfs = list(Path(tmp).glob("*.pdf"))
            if pdfs:
                pdf_path = str(pdfs[0])

    if not pdf_path or not os.path.exists(pdf_path):
        raise RuntimeError("LibreOffice PDF conversion failed.")

    # Use pdftotext (much faster than pypdf for large documents)
    try:
        res = subprocess.run(['pdftotext', '-layout', pdf_path, '-'], capture_output=True, text=True, timeout=120)
        pages = res.stdout.split('\f')
        if pages and not pages[-1].strip():
            pages = pages[:-1]
        result = [(i + 1, text) for i, text in enumerate(pages)]
    except (FileNotFoundError, subprocess.TimeoutExpired):
        # Fallback to pypdf
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        result = [(i + 1, page.extract_text() or "") for i, page in enumerate(reader.pages)]
    
    try:
        os.remove(pdf_path)
    except OSError:
        pass
    return result


def find_page_markers(doc):
    markers = []
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        if "Page Number" in style and para.text.strip():
            markers.append({"index": i, "text": para.text.strip()})
    return markers


def _strip_bidi(s):
    """Remove Unicode bidirectional control characters that pdftotext inserts."""
    return re.sub(r'[\u200e\u200f\u200b\u202a\u202b\u202c\u202d\u202e\u2066\u2067\u2068\u2069\ufeff]', '', s)


def find_marker_on_page(page_texts, marker_text, search_from_idx=0):
    """Search for marker_text in page_texts starting from search_from_idx.
    Returns (word_page_number, index_in_page_texts) or (None, None)."""
    # Pass 1: exact line match, searching forward only
    for idx in range(search_from_idx, len(page_texts)):
        wp, text = page_texts[idx]
        for line in text.split("\n"):
            cleaned = _strip_bidi(line).strip()
            if cleaned == marker_text:
                return wp, idx
    # Pass 2: word-boundary regex match on cleaned text, searching forward only
    # Skip matches preceded by "Page " (skipped-page messages in rendered PDF)
    pattern = r"(?:^|\s)" + re.escape(marker_text) + r"(?=\s|$)"
    for idx in range(search_from_idx, len(page_texts)):
        wp, text = page_texts[idx]
        cleaned = _strip_bidi(text)
        m = re.search(pattern, cleaned, re.MULTILINE)
        if m:
            # Check that this isn't inside a "Page X [Y]" or "[Page X skipped" context
            start = max(0, m.start() - 6)
            prefix = cleaned[start:m.start()+1].lower()
            if "page " in prefix:
                continue
            return wp, idx
    return None, None


def build_page_ranges(markers, page_texts):
    MAX_SPAN = 8  # No single source page should span more than this many Word pages
    # Build set of non-empty Word pages (empty pages are LibreOffice artifacts)
    nonempty_pages = {wp for wp, text in page_texts if text.strip()}
    results = []
    prev_wp = 0
    search_from_idx = 0
    for mi, m in enumerate(markers):
        wp, found_idx = find_marker_on_page(page_texts, m["text"], search_from_idx)
        if wp is None:
            wp = prev_wp + 1
            print(f"  WARNING: '{m['text']}' not found, guessing page {wp}")
        else:
            # Sanity check: if span would be too large, the match is probably wrong
            span = wp - prev_wp
            if span > MAX_SPAN and prev_wp > 0:
                print(f"  WARNING: '{m['text']}' found on page {wp} but span={span} exceeds max, guessing page {prev_wp+1}")
                wp = prev_wp + 1
                # Don't advance search_from_idx so next marker can still find its page
            else:
                search_from_idx = found_idx + 1
        pages = [p for p in range(prev_wp + 1, wp + 1) if p in nonempty_pages]
        if not pages:
            pages = [wp]  # Fallback: at least include the marker page
        results.append({
            "source": m["text"],
            "marker_index": m["index"],
            "word_pages": pages,
        })
        prev_wp = wp
    return results


def compute_labels(ranges):
    letters = "abcdefghijklmnopqrstuvwxyz"
    for r in ranges:
        if len(r["word_pages"]) <= 1:
            r["labels"] = [(r["word_pages"][0] if r["word_pages"] else 0, r["source"])]
        else:
            r["labels"] = []
            for i, wp in enumerate(r["word_pages"]):
                letter = letters[i] if i < len(letters) else str(i + 1)
                m = re.match(r"^(\d+)\s*\[(\d+)\]$", r["source"])
                if m:
                    r["labels"].append((wp, f"{m.group(1)}{letter} [{m.group(2)}{letter}]"))
                else:
                    r["labels"].append((wp, f"{r['source']}{letter}"))
    return ranges


def _strip_marks(s):
    """Strip Unicode combining marks (Hebrew vowels, cantillation) for fuzzy matching."""
    return ''.join(c for c in _strip_bidi(s) if not unicodedata.category(c).startswith('M'))


def build_paragraph_page_map(doc, page_texts):
    collapsed = [(wp, _strip_bidi(text).replace("\n", " ").replace("  ", " ")) for wp, text in page_texts]
    # Also build a marks-stripped version for Hebrew fallback matching
    collapsed_stripped = [(wp, _strip_marks(text).replace("\n", " ").replace("  ", " ")) for wp, text in page_texts]
    para_to_page = {}
    min_page = 1
    for i, para in enumerate(doc.paragraphs):
        pt = _strip_bidi(para.text.strip())
        if len(pt) < 5:
            continue
        sig = pt[:15]
        found = False
        for wp, ctext in collapsed:
            if wp < min_page:
                continue
            if sig in ctext:
                para_to_page[i] = wp
                min_page = wp
                found = True
                break
        if not found:
            # Fallback: strip combining marks (fixes Hebrew vowel mismatches
            # between docx and pdftotext output)
            sig_stripped = _strip_marks(pt[:15])
            if len(sig_stripped) >= 3:
                for wp, ctext in collapsed_stripped:
                    if wp < min_page:
                        continue
                    if sig_stripped in ctext:
                        para_to_page[i] = wp
                        min_page = wp
                        break
    return para_to_page, sorted(para_to_page.items())


def get_style(para):
    return para.style.name if para.style else ""


def is_new_reading(para):
    t = para.text.strip()
    return bool(re.match(r'^[A-Z]{2,}[\s,;:\.!]', t))


def is_attribution(para):
    """Is this paragraph an author attribution line?
    E.g., 'Sidney Greenberg (Adapted)', 'Chaim Stern', 'Psalms 42:2-6'"""
    s = get_style(para)
    # Style-based detection (most reliable)
    if "Attribution" in s:
        return True
    t = para.text.strip()
    if not t or len(t) > 100:
        return False
    # Only use text-based detection for unstyled or generic styles
    # Paragraphs with Translation, Hebrew, Section Header, etc. are NOT attributions
    if s and s not in ("Normal", "Body Text", "Default Paragraph Font", ""):
        return False
    # Short line with a name pattern: capitalized words, possibly with (Adapted), (Modified), etc.
    if re.match(r'^[A-Z][a-z]', t) and len(t) < 80:
        return True
    # Source citations like "Psalms 42:2-6", "Isaiah 6:3", "Talmud Berahot 34b"
    if re.match(r'^(Psalms?|Isaiah|Genesis|Exodus|Leviticus|Numbers|Deuteronomy|Proverbs|'
                r'Job|Ecclesiastes|Song of Songs|Lamentations|Daniel|Ezra|Nehemiah|'
                r'Chronicles|Samuel|Kings|Judges|Ruth|Esther|Amos|Hosea|Micah|'
                r'Jeremiah|Ezekiel|Habakkuk|Zechariah|Malachi|Talmud|Mishnah|Midrash|Zohar)', t):
        return True
    return False


def is_instruction_note(para):
    """Is this paragraph a navigation/instruction note?
    E.g., 'Additional meditations may be found on pages 1-20.'
          'Concluding prayers begin on page 1195.'
          'On Shabbat continue on page 132.'"""
    t = para.text.strip()
    if not t or len(t) > 200:
        return False
    t_lower = t.lower().lstrip('*')
    # Navigation/cross-reference notes
    if re.search(r'\b(continue|turn to|found on|begin on|see) (on )?page', t_lower):
        return True
    if re.search(r'\badditional .+ (may be |are )found\b', t_lower):
        return True
    if re.search(r'\bconcluding .+ begin\b', t_lower):
        return True
    return False


def is_keep_with_above(para):
    """Should this paragraph stay with the content above it?
    True for attributions and instruction notes."""
    return is_attribution(para) or is_instruction_note(para)


def is_section_header(para):
    s = get_style(para)
    return "Section Header" in s or "Heading" in s


def should_start_new_page(para):
    """Should this paragraph start a new page? (i.e., break BEFORE it)"""
    if is_section_header(para):
        return True
    if is_new_reading(para):
        return True
    return False


def adjust_for_short_title(insert_idx, needs_split, doc, mapped_list, wp):
    """Don't break after a short title-like paragraph — keep it with the next paragraph.
    A short (<60 char) paragraph followed by a longer paragraph of the same style
    is likely a title (e.g., a poem title before the poem body)."""
    if insert_idx is None or needs_split:
        return insert_idx
    para = doc.paragraphs[insert_idx]
    text = para.text.strip()
    if len(text) >= 60 or not text:
        return insert_idx
    style = get_style(para)
    if insert_idx + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[insert_idx + 1]
        next_style = get_style(next_para)
        if style == next_style and len(next_para.text.strip()) > len(text):
            page_paras = [pi for pi, wp2 in mapped_list if wp2 == wp]
            if insert_idx in page_paras:
                pos = page_paras.index(insert_idx)
                if pos > 0:
                    return page_paras[pos - 1]
    return insert_idx


def adjust_insert_for_grouping(insert_idx, needs_split, doc, mapped_list, wp):
    """Post-insert adjustments to prevent stranding headers and attributions.

    1. If the paragraph at insert_idx is a section header, move back so the
       header travels to the next page with its content.
    2. If the paragraph AFTER insert_idx is an attribution or instruction note,
       advance insert_idx to keep the attribution with its reading.
    """
    if insert_idx is None or needs_split:
        return insert_idx

    page_paras = [pi for pi, wp2 in mapped_list if wp2 == wp]

    # Rule 1: Don't strand a section header — push it to the next page
    para = doc.paragraphs[insert_idx]
    if is_section_header(para):
        pos = page_paras.index(insert_idx) if insert_idx in page_paras else -1
        if pos > 0:
            insert_idx = page_paras[pos - 1]
            # If the new candidate is also keep-with-above, back up further
            if is_keep_with_above(doc.paragraphs[insert_idx]) and pos > 1:
                insert_idx = page_paras[pos - 2]

    # Rule 2: Don't separate an attribution from its reading
    # Check the next paragraph in the DOCUMENT (not just on the page)
    if insert_idx + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[insert_idx + 1]
        if is_keep_with_above(next_para):
            # Include the attribution and any trailing instruction notes
            candidate = insert_idx + 1
            while candidate + 1 < len(doc.paragraphs):
                following = doc.paragraphs[candidate + 1]
                if is_keep_with_above(following):
                    candidate += 1
                else:
                    break
            # Only advance if this doesn't consume the page number marker
            next_style = get_style(doc.paragraphs[candidate])
            if "Page Number" not in next_style:
                insert_idx = candidate

    return insert_idx


def find_best_insert_point(wp, mapped_list, para_to_page, doc):
    """
    Find the best paragraph to insert a page number AFTER.

    Returns (para_index, split_flag, structural).
    structural=True means a meaningful boundary was found (header, reading,
    attribution, style transition). structural=False means the default
    fallback was used.

    Rules (in priority order):
    1. Section headers ALWAYS go to the next page — break before them
       regardless of position on the page.
    2. New readings (ALL CAPS start, excl. headers) in the bottom 2/3 go to next page.
    3. Hebrew→Transliteration boundary in the bottom 2/3: split transliteration.
    4. Attribution lines (with trailing notes) in the bottom 2/3 mark
       a natural reading boundary — break after them.
    5. Default: scan for ANY attribution on the page; break after it.
    6. Natural paragraph breaks: period-ending paragraphs in the bottom 2/3.
    7. Fallback: second-to-last paragraph.
    8. If only 1 paragraph: split it.
    9. After choosing, check if only 1 paragraph would remain after the break.
       If that paragraph is long (>200 chars), split it instead of leaving it alone.
    """
    page_paras = [pi for pi, wp2 in mapped_list if wp2 == wp]

    if len(page_paras) == 0:
        return None, False, False

    if len(page_paras) == 1:
        return page_paras[0], "split_solo", True

    # Helper: given a position in page_paras, skip past any trailing
    # attribution + instruction note paragraphs so they stay with the reading.
    def skip_trailing_keep_withs(pos):
        """Advance pos past consecutive attribution/instruction-note paragraphs."""
        while pos + 1 < len(page_paras):
            next_p = doc.paragraphs[page_paras[pos + 1]]
            if is_keep_with_above(next_p):
                pos += 1
            else:
                break
        return pos

    # Pass 1: Scan ALL paragraphs for section headers (always push to next page)
    for i in range(len(page_paras) - 1, 0, -1):
        para = doc.paragraphs[page_paras[i]]
        if is_section_header(para):
            candidate_idx = i - 1
            # Don't strand another header — keep scanning back
            while candidate_idx > 1 and is_section_header(doc.paragraphs[page_paras[candidate_idx]]):
                candidate_idx -= 1
            if is_section_header(doc.paragraphs[page_paras[candidate_idx]]):
                continue
            # Check if the paragraph before the candidate is an attribution
            # or instruction note — if so, break after that instead.
            # This keeps the attribution with its reading on the current page
            # and pushes the new content (candidate onward) to the next page.
            if candidate_idx > 0:
                prev_p = doc.paragraphs[page_paras[candidate_idx - 1]]
                if is_keep_with_above(prev_p):
                    candidate_idx -= 1
            return page_paras[candidate_idx], False, True

    # Pass 2: Scan bottom two-thirds for new readings, style transitions
    midpoint = max(len(page_paras) // 3, 1)

    for i in range(len(page_paras) - 1, midpoint - 1, -1):
        para = doc.paragraphs[page_paras[i]]
        prev_para = doc.paragraphs[page_paras[i - 1]]
        para_style = get_style(para)
        prev_style = get_style(prev_para)

        if is_new_reading(para) and not is_section_header(para):
            # Scan backward past consecutive new readings
            j = i
            while j > 1 and should_start_new_page(doc.paragraphs[page_paras[j - 1]]):
                j -= 1
            # But don't go above midpoint
            if j >= midpoint:
                result_idx = page_paras[j - 1]
                # Don't return a section header as the break point
                if not is_section_header(doc.paragraphs[result_idx]):
                    return result_idx, False, True
            else:
                result_idx = page_paras[i - 1]
                if not is_section_header(doc.paragraphs[result_idx]):
                    return result_idx, False, True

        # Hebrew → Transliteration boundary: try to keep transliteration with Hebrew
        # by splitting the transliteration at a period rather than pushing it all away
        if "Transliteration" in para_style and "Hebrew" in prev_style:
            # Check if this transliteration has periods we can split at
            if '.' in para.text:
                return page_paras[i], "split_translit", True
            # No periods — fall back to breaking before transliteration
            return page_paras[i - 1], False, True

        # Any Hebrew after non-Hebrew
        if "Hebrew" in para_style and "Hebrew" not in prev_style:
            return page_paras[i - 1], False, True

    # Pass 3: Scan bottom two-thirds for attributions as natural break points
    for i in range(len(page_paras) - 1, midpoint - 1, -1):
        para = doc.paragraphs[page_paras[i]]
        if is_attribution(para):
            # Found an attribution — skip past any trailing instruction notes
            pos_in_page = i
            pos_in_page = skip_trailing_keep_withs(pos_in_page)
            # Don't pick the very last paragraph (nothing would go to next page)
            if pos_in_page < len(page_paras) - 1:
                return page_paras[pos_in_page], False, True

    # Default: scan entire page for any attribution (even in top half)
    for i in range(len(page_paras) - 1, 0, -1):
        para = doc.paragraphs[page_paras[i]]
        if is_attribution(para):
            pos_in_page = i
            pos_in_page = skip_trailing_keep_withs(pos_in_page)
            if pos_in_page < len(page_paras) - 1:
                return page_paras[pos_in_page], False, True

    # Pass 4: Look for natural stanza/paragraph breaks in lower two-thirds.
    # A paragraph ending with terminal punctuation (.!?) followed by another
    # paragraph of the same style that starts a new thought.
    for i in range(len(page_paras) - 1, midpoint - 1, -1):
        para = doc.paragraphs[page_paras[i]]
        # This paragraph should not be keep-with-above and not be the last
        if i >= len(page_paras) - 1:
            continue
        next_para = doc.paragraphs[page_paras[i + 1]] if i + 1 < len(page_paras) else None
        if next_para and is_keep_with_above(next_para):
            continue
        t = para.text.rstrip()
        if t and t[-1] in '.!?':
            return page_paras[i], False, True

    # Last resort: second-to-last paragraph
    best = page_paras[-2]

    # If the last paragraph is keep-with-above (attribution/note), don't
    # split it from the content before it — use last instead
    last_para = doc.paragraphs[page_paras[-1]]
    if is_keep_with_above(last_para):
        best = page_paras[-1]

    # Don't break right before an attribution or instruction note —
    # keep it with the reading above
    best_pos = page_paras.index(best) if best in page_paras else -1
    if best_pos >= 0 and best_pos + 1 < len(page_paras):
        next_para = doc.paragraphs[page_paras[best_pos + 1]]
        if is_keep_with_above(next_para) and best_pos > 0:
            best = page_paras[best_pos - 1]

    # Check: would only 1 long paragraph remain after the break?
    remaining_after = page_paras[page_paras.index(best) + 1:] if best in page_paras else []
    if len(remaining_after) == 1:
        remaining_para = doc.paragraphs[remaining_after[0]]
        remaining_style = get_style(remaining_para)
        # If it's a long translation or transliteration, offer to split it
        if len(remaining_para.text) > 200 and ("Translation" in remaining_style or "Transliteration" in remaining_style):
            return remaining_after[0], True, False

    return best, False, False


def split_paragraph_at_last_period(doc, para_index):
    """Split a paragraph at the LAST sentence boundary, keeping as much as
    possible on the current page (with the preceding Hebrew).
    Returns para_index on success, None on failure."""
    para = doc.paragraphs[para_index]
    text = para.text
    # Find all period-based sentence boundaries
    sentence_ends = []
    for m in re.finditer(r'[.]\s+', text):
        sentence_ends.append(m.start() + 1)
    if not sentence_ends:
        return None

    # Try from last period backwards — skip any that give an empty second half
    best = None
    for candidate in reversed(sentence_ends):
        first_half = text[:candidate].rstrip()
        second_half = text[candidate:].lstrip()
        if first_half and second_half and len(second_half) >= 10:
            best = candidate
            break
    if best is None:
        return None
    first_half = text[:best].rstrip()
    second_half = text[best:].lstrip()
    if not first_half or not second_half:
        return None

    para_element = para._element
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = first_half
    else:
        para.add_run(first_half)

    new_p = deepcopy(para_element)
    for r_elem in new_p.findall(qn('w:r')):
        new_p.remove(r_elem)
    r = etree.SubElement(new_p, qn('w:r'))
    orig_runs = para_element.findall(qn('w:r'))
    if orig_runs:
        orig_rPr = orig_runs[0].find(qn('w:rPr'))
        if orig_rPr is not None:
            r.insert(0, deepcopy(orig_rPr))
    t = etree.SubElement(r, qn('w:t'))
    t.text = second_half
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    para_element.addnext(new_p)
    return para_index


def split_paragraph_at_blank_line(doc, para_index):
    """Split a paragraph at the FIRST blank-line boundary (\\n\\n).
    Used for transliteration paragraphs that contain multiple logical sections.
    Keeps the opening section with the preceding Hebrew."""
    para = doc.paragraphs[para_index]
    text = para.text
    # Find all blank-line boundaries
    boundaries = [m.start() for m in re.finditer(r'\n\n', text)]
    if not boundaries:
        return None

    # Pick the FIRST blank-line boundary to keep the opening with Hebrew
    best = boundaries[0]
    first_half = text[:best].rstrip()
    second_half = text[best:].lstrip()
    if not first_half or not second_half or len(second_half) < 10:
        return None

    para_element = para._element
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = first_half
    else:
        para.add_run(first_half)

    new_p = deepcopy(para_element)
    for r_elem in new_p.findall(qn('w:r')):
        new_p.remove(r_elem)
    r = etree.SubElement(new_p, qn('w:r'))
    orig_runs = para_element.findall(qn('w:r'))
    if orig_runs:
        orig_rPr = orig_runs[0].find(qn('w:rPr'))
        if orig_rPr is not None:
            r.insert(0, deepcopy(orig_rPr))
    t = etree.SubElement(r, qn('w:t'))
    t.text = second_half
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    para_element.addnext(new_p)
    return para_index


def split_paragraph_at_midpoint(doc, para_index, target_pct=0.65):
    para = doc.paragraphs[para_index]
    text = para.text
    sentence_ends = []
    for m in re.finditer(r'[.!?]\s+(?=[A-Z"\'])', text):
        sentence_ends.append(m.start() + 1)
    if not sentence_ends:
        for m in re.finditer(r'[.!?]\s+', text):
            sentence_ends.append(m.start() + 1)
    if not sentence_ends:
        for m in re.finditer(r'[,;]\s+', text):
            sentence_ends.append(m.start() + 1)
    if not sentence_ends:
        target_fallback = int(len(text) * target_pct)
        sp = text.rfind(' ', 0, target_fallback + 20)
        if sp > 10:
            sentence_ends = [sp]
    if not sentence_ends:
        return None

    # Bias toward a later split so more content stays on the current page.
    # Default 65%; single-paragraph pages use 75% since the page number
    # marker above already occupies space.
    target = int(len(text) * target_pct)
    best = min(sentence_ends, key=lambda pos: abs(pos - target))
    first_half = text[:best].rstrip()
    second_half = text[best:].lstrip()
    if not first_half or not second_half:
        return None

    para_element = para._element
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = first_half
    else:
        para.add_run(first_half)

    new_p = deepcopy(para_element)
    for r_elem in new_p.findall(qn('w:r')):
        new_p.remove(r_elem)
    r = etree.SubElement(new_p, qn('w:r'))
    orig_runs = para_element.findall(qn('w:r'))
    if orig_runs:
        orig_rPr = orig_runs[0].find(qn('w:rPr'))
        if orig_rPr is not None:
            r.insert(0, deepcopy(orig_rPr))
    t = etree.SubElement(r, qn('w:t'))
    t.text = second_half
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    para_element.addnext(new_p)
    return para_index


def make_page_num_xml(label, style_id="PageNumber"):
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    # pStyle must come first in pPr per OOXML schema
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), style_id)
    spacing = etree.SubElement(pPr, qn('w:spacing'))
    spacing.set(qn('w:before'), '400')
    spacing.set(qn('w:after'), '0')
    jc = etree.SubElement(pPr, qn('w:jc'))
    jc.set(qn('w:val'), 'center')

    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Tahoma')
    rFonts.set(qn('w:hAnsi'), 'Tahoma')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '44')
    color = etree.SubElement(rPr, qn('w:color'))
    color.set(qn('w:val'), '000000')
    t = etree.SubElement(r, qn('w:t'))
    t.text = label
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    r2 = etree.SubElement(p, qn('w:r'))
    br = etree.SubElement(r2, qn('w:br'))
    br.set(qn('w:type'), 'page')
    return p


def _has_page_break(para):
    """Check if a paragraph contains a w:br type='page'."""
    for run_el in para._element.findall(qn('w:r')):
        for br in run_el.findall(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
    return False


def _add_page_break(para):
    """Add a page break run to a paragraph element."""
    r = etree.SubElement(para._element, qn('w:r'))
    br = etree.SubElement(r, qn('w:br'))
    br.set(qn('w:type'), 'page')


def _remove_paragraph(para):
    """Remove a paragraph element from its parent."""
    el = para._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _get_base_and_suffix(label):
    """Parse '1195b' -> ('1195', 'b') or '1197' -> ('1197', '').
    Also handles dual labels like '1195b [42b]'."""
    # Strip any bracketed dual label
    base_part = label.split('[')[0].strip()
    m = re.match(r'^(\d+)([a-z]?)$', base_part)
    if m:
        return m.group(1), m.group(2)
    return base_part, ''


def _next_suffix(suffix):
    """Return the next letter suffix: '' -> 'a', 'a' -> 'b', 'b' -> 'c', etc."""
    if not suffix:
        return 'a'
    return chr(ord(suffix) + 1)


def _rebuild_label(base, suffix, original_label):
    """Rebuild a label with a new suffix, preserving dual-label format if present."""
    bracket = ''
    m = re.search(r'\[(\d+)[a-z]?\]', original_label)
    if m:
        bracket = f' [{m.group(1)}{suffix}]' if suffix else f' [{m.group(1)}]'
    return f'{base}{suffix}{bracket}'


def cleanup_pass(doc, pn_style_id="PageNumber"):
    """Post-processing cleanup pass to tidy page number structure.

    Three rules applied in order:

    Rule 1 — Collapse empty suffix pairs:
        When two Page Number paragraphs appear consecutively (e.g., 1197a
        followed immediately by 1197b) with no content between them, the
        second page is empty.  Remove it (and its trailing blank separator)
        and strip the letter suffix from the first (1197a → 1197).

    Rule 2 — Label unlabeled overflow pages:
        When an overflow page number (has a page break, starts a new page)
        is followed by content, then an empty paragraph with a page break,
        then MORE content before the next page number — that second run of
        content is on a page with no label.  Insert the next sequential
        letter suffix (e.g., 1195b → add 1195c).

    Rule 3 — Consolidate blank page separators:
        When a page number paragraph (no page break) is immediately followed
        by an empty paragraph whose only purpose is a page break, move the
        break onto the page number and delete the empty paragraph.  This
        prevents stray blank pages in the rendered output.
    """
    print("Step 7: Cleanup pass...")
    stats = {"collapsed": 0, "labeled": 0, "consolidated": 0}

    # Work entirely with raw lxml elements to avoid python-docx wrapper
    # issues on paragraphs inserted as raw XML by Step 6.
    body = doc.element.body
    w_p_tag = qn('w:p')
    w_pStyle_tag = qn('w:pStyle')
    w_pPr_tag = qn('w:pPr')
    w_r_tag = qn('w:r')
    w_br_tag = qn('w:br')
    w_t_tag = qn('w:t')
    w_type_attr = qn('w:type')
    w_val_attr = qn('w:val')

    def _el_style(el):
        pPr = el.find(w_pPr_tag)
        if pPr is not None:
            pStyle = pPr.find(w_pStyle_tag)
            if pStyle is not None:
                return pStyle.get(w_val_attr, "")
        return ""

    def _el_text(el):
        parts = []
        for r in el.findall(w_r_tag):
            for t in r.findall(w_t_tag):
                if t.text:
                    parts.append(t.text)
        return ''.join(parts)

    def _el_has_page_break(el):
        for r in el.findall(w_r_tag):
            for br in r.findall(w_br_tag):
                if br.get(w_type_attr) == 'page':
                    return True
        return False

    def _el_add_page_break(el):
        r = etree.SubElement(el, w_r_tag)
        br = etree.SubElement(r, w_br_tag)
        br.set(w_type_attr, 'page')

    def _el_set_text(el, new_text):
        """Set text on a paragraph, preserving page break runs."""
        first_text_set = False
        for run_el in el.findall(w_r_tag):
            t_el = run_el.find(w_t_tag)
            br_el = run_el.find(w_br_tag)
            if br_el is not None:
                continue
            if t_el is not None:
                if not first_text_set:
                    t_el.text = new_text
                    first_text_set = True
                else:
                    t_el.text = ""

    def _get_paragraphs():
        return [el for el in body if el.tag == w_p_tag]

    # ── Rule 1: Collapse empty suffix pairs ────────────────────────────
    paragraphs = _get_paragraphs()
    i = len(paragraphs) - 2
    while i >= 0:
        el = paragraphs[i]
        nxt_el = paragraphs[i + 1]
        s1 = _el_style(el)
        s2 = _el_style(nxt_el)

        if "PageNumber" in s1 and "PageNumber" in s2:
            t1 = _el_text(el).strip()
            t2 = _el_text(nxt_el).strip()
            base1, suf1 = _get_base_and_suffix(t1)
            base2, suf2 = _get_base_and_suffix(t2)

            if base1 == base2 and suf1 and suf2:
                # Remove the trailing empty paragraph (if present after Xb)
                if i + 2 < len(paragraphs):
                    trailing = paragraphs[i + 2]
                    if not _el_text(trailing).strip() and "PageNumber" not in _el_style(trailing):
                        body.remove(trailing)

                # Remove the Xb page number
                body.remove(nxt_el)

                # Rename Xa → X
                new_label = _rebuild_label(base1, '', t1)
                _el_set_text(el, new_label)

                print(f"    Collapsed: '{t1}' + '{t2}' → '{new_label}'")
                stats["collapsed"] += 1
                # Refresh list
                paragraphs = _get_paragraphs()
        i -= 1

    # ── Rule 2: Label unlabeled overflow pages ─────────────────────────
    paragraphs = _get_paragraphs()
    page_nums = []
    for i, el in enumerate(paragraphs):
        if "PageNumber" in _el_style(el):
            page_nums.append((i, _el_text(el).strip(), _el_has_page_break(el)))

    insertions = []  # (empty_para_element, new_label)
    for pi in range(len(page_nums) - 1):
        idx, label, brk = page_nums[pi]
        next_idx, next_label, next_brk = page_nums[pi + 1]

        if not brk:
            continue
        base, suffix = _get_base_and_suffix(label)
        if not suffix:
            continue

        for j in range(idx + 1, next_idx):
            el = paragraphs[j]
            if not _el_text(el).strip() and _el_has_page_break(el) and "PageNumber" not in _el_style(el):
                has_content_after = any(
                    _el_text(paragraphs[k]).strip()
                    for k in range(j + 1, next_idx)
                )
                if has_content_after:
                    new_suffix = _next_suffix(suffix)
                    new_label = _rebuild_label(base, new_suffix, label)
                    insertions.append((el, new_label))
                    break

    for empty_el, new_label in reversed(insertions):
        new_pn = make_page_num_xml(new_label, pn_style_id)
        body.replace(empty_el, new_pn)
        print(f"    Labeled: inserted '{new_label}' (unlabeled overflow page)")
        stats["labeled"] += 1

    # ── Rule 3: Consolidate blank page separators ──────────────────────
    paragraphs = _get_paragraphs()
    i = len(paragraphs) - 2
    while i >= 0:
        el = paragraphs[i]
        style = _el_style(el)

        if "PageNumber" in style and not _el_has_page_break(el):
            nxt_el = paragraphs[i + 1]
            nxt_style = _el_style(nxt_el)
            nxt_text = _el_text(nxt_el).strip()
            if not nxt_text and "PageNumber" not in nxt_style and _el_has_page_break(nxt_el):
                _el_add_page_break(el)
                body.remove(nxt_el)
                stats["consolidated"] += 1
        i -= 1

    print(f"    Collapsed {stats['collapsed']} empty suffix pair(s).")
    print(f"    Labeled {stats['labeled']} unlabeled overflow page(s).")
    print(f"    Consolidated {stats['consolidated']} blank separator(s).\n")
    return stats


def process(input_path, output_path, lo_bin):
    print("Step 1: Analyzing document...")
    doc = Document(str(input_path))
    markers = find_page_markers(doc)
    print(f"  {len(markers)} source page marker(s).\n")

    # Detect the page number style ID from existing markers
    pn_style_id = "PageNumber"
    if markers:
        para = doc.paragraphs[markers[0]["index"]]
        if para.style and para.style.style_id:
            pn_style_id = para.style.style_id
    print(f"  Page number style: {pn_style_id}")

    print("Step 2: Rendering with LibreOffice...")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name); tmp_path = tmp.name
    try:
        page_texts = render_to_pdf_pages(tmp_path, lo_bin)
    finally:
        os.unlink(tmp_path)
    total_pages = len(page_texts)
    print(f"  {total_pages} Word page(s).\n")

    print("Step 3: Mapping source pages to Word pages...")
    ranges = build_page_ranges(markers, page_texts)
    overflow = [r for r in ranges if len(r["word_pages"]) > 1]
    print(f"  {len(ranges) - len(overflow)} single-page, {len(overflow)} overflow.\n")

    print("Step 4: Computing labels...")
    ranges = compute_labels(ranges)
    for r in ranges:
        if len(r["labels"]) > 1:
            print(f"    {r['source']} → {', '.join(l for _, l in r['labels'])}")
    print()

    print("Step 5: Updating and repositioning markers...")
    updated = 0
    repositioned = 0

    if overflow:
        # Build paragraph-page map now since we need it for repositioning
        para_to_page, mapped_list = build_paragraph_page_map(doc, page_texts)

    for r in ranges:
        last_label = r["labels"][-1][1]
        marker_para = doc.paragraphs[r["marker_index"]]

        # Update text if label changed
        if last_label != r["source"]:
            for run in marker_para.runs:
                run.text = ""
            if marker_para.runs:
                marker_para.runs[0].text = last_label
            else:
                run = marker_para.add_run(last_label)
                run.font.name = "Tahoma"
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            updated += 1

        # For overflow sources, check if the marker should be repositioned
        # on its page (move it before headers/new readings)
        if len(r["labels"]) > 1:
            marker_wp = r["labels"][-1][0]  # Word page the marker is on
            page_paras = [pi for pi, wp2 in mapped_list if wp2 == marker_wp]

            if len(page_paras) >= 2:
                # Find best position on this page using same logic
                best_idx, _, is_structural = find_best_insert_point(marker_wp, mapped_list, para_to_page, doc)

                # Only reposition if there's a structural boundary (header,
                # new reading, attribution). Skip for default fallbacks.
                if best_idx is not None and best_idx < r["marker_index"] and is_structural:
                        # Determine: should we insert before or after best_idx?
                        # If best_idx itself is a new reading or header, go before it
                        best_para = doc.paragraphs[best_idx]
                        insert_before_best = should_start_new_page(best_para)

                        marker_element = marker_para._element
                        parent = marker_element.getparent()
                        parent.remove(marker_element)

                        # Add page break to marker if needed
                        has_break = False
                        for run_el in marker_element.findall(qn('w:r')):
                            for br in run_el.findall(qn('w:br')):
                                if br.get(qn('w:type')) == 'page':
                                    has_break = True
                        if not has_break:
                            r2 = etree.SubElement(marker_element, qn('w:r'))
                            br = etree.SubElement(r2, qn('w:br'))
                            br.set(qn('w:type'), 'page')

                        ref = doc.paragraphs[best_idx]._element
                        if insert_before_best:
                            ref.addprevious(marker_element)
                        else:
                            ref.addnext(marker_element)
                        repositioned += 1

                        next_content = ""
                        scan_start = best_idx + (0 if insert_before_best else 1)
                        for j in range(scan_start, min(scan_start + 3, len(doc.paragraphs))):
                            t = doc.paragraphs[j].text.strip()
                            if t and "Page Number" not in get_style(doc.paragraphs[j]):
                                next_content = t[:40]
                                break
                        pos = "before" if insert_before_best else "after"
                        print(f"    Moved '{last_label}' {pos} [{best_idx}] '{next_content}...'")

    print(f"  Updated {updated} marker(s), repositioned {repositioned}.\n")

    if overflow:
        print("Step 6: Finding insertion points...")
        # para_to_page and mapped_list already built in Step 5

        actions = []
        for r in ranges:
            if len(r["labels"]) <= 1:
                continue
            for wp, label in r["labels"][:-1]:
                insert_idx, needs_split, _ = find_best_insert_point(wp, mapped_list, para_to_page, doc)
                insert_idx = adjust_for_short_title(insert_idx, needs_split, doc, mapped_list, wp)
                insert_idx = adjust_insert_for_grouping(insert_idx, needs_split, doc, mapped_list, wp)
                if insert_idx is None:
                    print(f"    WARNING: No point for '{label}' (page {wp})")
                    continue

                next_content = ""
                for j in range(insert_idx + 1, min(insert_idx + 3, len(doc.paragraphs))):
                    t = doc.paragraphs[j].text.strip()
                    if t:
                        s = get_style(doc.paragraphs[j])
                        next_content = f"[{s}] {t[:35]}"
                        break

                if needs_split == "split_translit":
                    actions.append(("split_translit", insert_idx, label))
                    print(f"    Split translit [{insert_idx}] for '{label}' (keep max with Hebrew)")
                elif needs_split == "split_solo":
                    actions.append(("split_solo", insert_idx, label))
                    print(f"    Split solo [{insert_idx}] for '{label}' (single-para page, 75% bias)")
                elif needs_split:
                    actions.append(("split", insert_idx, label))
                    print(f"    Split [{insert_idx}] for '{label}' (before '{next_content}...')")
                else:
                    print(f"    '{label}' after [{insert_idx}] (before '{next_content}...')")
                    actions.append(("insert", insert_idx, label))

        actions.sort(key=lambda x: x[1], reverse=True)

        inserted = 0
        splits = 0
        for action in actions:
            atype, para_idx, label = action
            if atype == "split_translit":
                para_text = doc.paragraphs[para_idx].text
                # Don't try to split short transliteration paragraphs
                if len(para_text) < 150:
                    # Too short to split — insert before it instead
                    ref = doc.paragraphs[para_idx]._element
                    new_p = make_page_num_xml(label, pn_style_id)
                    ref.addprevious(new_p)
                    inserted += 1
                else:
                    # Strategy 1: If transliteration has blank-line paragraph breaks,
                    # split at the first one (keeps opening section with Hebrew)
                    result = split_paragraph_at_blank_line(doc, para_idx)
                    if result is not None:
                        splits += 1
                        ref = doc.paragraphs[para_idx]._element
                        new_p = make_page_num_xml(label, pn_style_id)
                        ref.addnext(new_p)
                        inserted += 1
                    else:
                        # Strategy 2: Split at a period near 40% (balanced split)
                        result = split_paragraph_at_midpoint(doc, para_idx, target_pct=0.40)
                        if result is not None:
                            splits += 1
                            ref = doc.paragraphs[para_idx]._element
                            new_p = make_page_num_xml(label, pn_style_id)
                            ref.addnext(new_p)
                            inserted += 1
                        else:
                            # Strategy 3: Fall back to inserting before transliteration
                            ref = doc.paragraphs[para_idx]._element
                            new_p = make_page_num_xml(label, pn_style_id)
                            ref.addprevious(new_p)
                            inserted += 1
            elif atype == "split_solo":
                result = split_paragraph_at_midpoint(doc, para_idx, target_pct=0.75)
                if result is not None:
                    splits += 1
                ref = doc.paragraphs[para_idx]._element
                new_p = make_page_num_xml(label, pn_style_id)
                ref.addnext(new_p)
                inserted += 1
            elif atype == "split":
                result = split_paragraph_at_midpoint(doc, para_idx)
                if result is not None:
                    splits += 1
                ref = doc.paragraphs[para_idx]._element
                new_p = make_page_num_xml(label, pn_style_id)
                ref.addnext(new_p)
                inserted += 1
            else:
                ref = doc.paragraphs[para_idx]._element
                new_p = make_page_num_xml(label, pn_style_id)
                ref.addnext(new_p)
                inserted += 1

        print(f"\n  Inserted {inserted}, split {splits}.")
    else:
        print("Step 6: No overflow.")

    cleanup_pass(doc, pn_style_id)

    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    return total_pages, ranges


def main():
    parser = argparse.ArgumentParser(description="Add page numbers to Siddur Processor output.")
    parser.add_argument("docx", help="Path to .docx from Siddur Processor")
    parser.add_argument("--output", default=None, help="Output path (default: _numbered suffix)")
    args = parser.parse_args()

    path = Path(args.docx)
    if not path.exists():
        print(f"ERROR: {path} not found.", file=sys.stderr); sys.exit(1)

    output = args.output
    if not output:
        stem = str(path.with_suffix(""))
        # Extract version number from _processed-N / _processed_N patterns
        m = re.search(r'_processed[-_]?(\d+)?$', stem)
        if m:
            base = stem[:m.start()]
            version = m.group(1)
        else:
            base = stem
            version = None
        version_tag = f"_v{version}" if version else ""
        output = f"{base}_numbered{version_tag}.docx"
    lo = find_libreoffice()
    if not lo:
        print("ERROR: LibreOffice not found.", file=sys.stderr); sys.exit(1)

    print(f"LibreOffice: {lo}\n")
    print(f"Siddur Page Numberer")
    print(f"{'=' * 50}")
    print(f"Input:  {path.name}")
    print(f"Output: {Path(output).name}\n")

    total, ranges = process(str(path), output, lo)
    ovf = sum(1 for r in ranges if len(r["labels"]) > 1)

    print(f"\n{'=' * 50}")
    print(f"DONE!")
    print(f"  Word pages: {total}")
    if ovf:
        print(f"  Overflow:   {ovf} source page(s) got letter suffixes")
    print(f"  Output:     {output}")
    print(f"{'=' * 50}")


if __name__ == "__main__":
    main()
